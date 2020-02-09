# -*- coding: utf-8 -*-
"""
@Author     :
@Date       :2020/2/9 16:58
@Email      :
@File       :taotiehai.py
@Description:
@Software   :PyCharm
"""

import time
from util import util
import re
from docx import Document


def remove_tag(html):
    dr = re.compile(r'<[^>]+>', re.S)
    return dr.sub('', html)


def main():
    # https://xueqiu.com/v4/statuses/user_timeline.json?page=2&user_id=1314783718&type=4&_=1581238460615
    url = 'https://xueqiu.com/v4/statuses/user_timeline.json'
    header = {'Accept': '*/*',
              'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.86 Safari/537.36'}
    param = {'page': 1,
             'user_id': '1314783718',
             'type': '4',
             '_': int(round(time.time()*1000))}
    doc = Document()
    for i in range(1, 75):  # 目前只有74页 偷懒手动写死
        param['page'] = i
        quote = util.get_json_by_get(url, param, header, util.get_cookie_from_chrome('.xueqiu.com'))['statuses']
        for ans in quote:
            if ans['created_at'] > time.mktime(time.strptime('2019-01-01', '%Y-%m-%d'))*1000:
                if remove_tag(ans['retweeted_status']['text']).find('转债') < 0:
                    doc.add_paragraph('***')
                doc.add_paragraph(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(float(ans['retweeted_status']['created_at'])/1000)))
                doc.add_paragraph('@' + ans['retweeted_status']['user']['screen_name'] + ':' + remove_tag(ans['retweeted_status']['text']))
                doc.add_paragraph(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(float(ans['created_at'])/1000)))
                doc.add_paragraph('@饕餮海:' + remove_tag(ans['text']))
                doc.add_paragraph('#'*20)
    doc.save('D:/taotiehai.docx')


if __name__ == '__main__':
    main()
